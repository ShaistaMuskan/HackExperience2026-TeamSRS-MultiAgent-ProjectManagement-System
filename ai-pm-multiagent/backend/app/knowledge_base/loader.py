"""
Shared seed-knowledge-base loader.

Used by both MockSearchClient (in-memory demo mode, app/integrations/mock/
mock_search_client.py) and scripts/ingest_knowledge_base.py (real Azure AI
Search ingestion) - one place to add support for new source file types so
both paths stay in sync and never drift apart.

Supported file types:
  .md    - read as plain text.
  .docx  - paragraph + table text extracted via python-docx (pip install python-docx).

Category is derived from the filename's `category__slug` prefix convention
(e.g. methodology__agile_scrum.md -> category "methodology"). Files that
don't follow that convention - e.g. reference docs someone drops into
seed_docs/, like the RelAI_Knowledge_Base_Part_*.docx files - fall back to
category "knowledge_base" rather than being skipped.
"""
from __future__ import annotations
import re
from pathlib import Path
from typing import Any

SEED_DOCS_DIR = Path(__file__).resolve().parent / "seed_docs"

_SUPPORTED_SUFFIXES = (".md", ".docx")


def _slugify(name: str) -> str:
    return re.sub(r"[^a-z0-9]+", "_", name.lower()).strip("_")


def _read_docx(path: Path) -> str:
    from docx import Document  # python-docx

    doc = Document(str(path))
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def load_seed_docs() -> list[dict[str, Any]]:
    """Returns [{"id", "content", "metadata": {"source", "filename", "category"}}, ...]."""
    if not SEED_DOCS_DIR.exists():
        return []

    out: list[dict[str, Any]] = []
    for path in sorted(SEED_DOCS_DIR.glob("*")):
        suffix = path.suffix.lower()
        if suffix not in _SUPPORTED_SUFFIXES:
            continue

        try:
            content = path.read_text(encoding="utf-8") if suffix == ".md" else _read_docx(path)
        except ImportError:
            # python-docx not installed - skip .docx files rather than crashing
            # the whole load (the .md seed docs still work fine without it).
            continue
        except Exception:
            continue

        if not content.strip():
            continue

        stem = path.stem
        category = stem.split("__", 1)[0] if "__" in stem else "knowledge_base"
        out.append({
            "id": _slugify(stem),
            "content": content,
            "metadata": {"source": "seed_kb", "filename": path.name, "category": category},
        })
    return out
